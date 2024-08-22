from docx import Document


def extract_hyperlinks(document):
    hyperlinks = {}

    for paragraph in document.paragraphs:
        for link in paragraph.hyperlinks:
            hyperlinks[link.text] = link.url
    return hyperlinks


def replace_quotes(text):
    # Replace English single quotes with Chinese single ending quotes
    text = text.replace("'", "’")

    # Replace English double quotes with Chinese double quotes
    parts = text.split('"')
    for i in range(len(parts)):
        if i % 2 == 0:
            # Even indices: outside of quotes, do nothing
            continue
        else:
            # Odd indices: inside of quotes, replace with Chinese quotes
            parts[i] = '“' + parts[i] + '”'
    return ''.join(parts)


def parse_word_document(file_path):
    doc = Document(file_path)
    subject = ""
    main_news = []
    also_featured_news = []
    current_section = None
    hyperlinks = extract_hyperlinks(doc)

    paragraphs = iter(doc.paragraphs)  # Create an iterator from the list of paragraphs

    for para in paragraphs:
        text = replace_quotes(para.text.strip())

        if text.startswith("Email subject line"):
            subject = text.split(":", 1)[1].strip()
            continue

        if "Also featured" in text:
            current_section = "Also_Featured"
            continue

        if current_section is None:
            current_section = "Main"

        if text:
            if current_section == "Main" or current_section == "Also_Featured":
                title = text
                content_paragraph = next(paragraphs)
                content_text = replace_quotes(content_paragraph.text.strip())

                # Initialize text parts for link separation
                text_before_link = ""
                text_to_link = ""
                content_link = ""
                text_after_link = ""

                # Check if the content has a hyperlink and separate accordingly
                for linked_text, link in hyperlinks.items():
                    if linked_text in content_text:
                        text_parts = content_text.split(linked_text)
                        text_before_link = text_parts[0]
                        text_to_link = linked_text
                        content_link = link
                        text_after_link = text_parts[1] if len(text_parts) > 1 else ""
                        break
                else:
                    text_before_link = content_text

                news_item = {
                    "Section": current_section,
                    "Title": title,
                    "Content_TextBeforeLink": text_before_link,
                    "Content_TextToLink": text_to_link,
                    "Content_Link": content_link,
                    "Content_TextAfterLink": text_after_link
                }

                if current_section == "Main":
                    main_news.append(news_item)
                else:
                    also_featured_news.append(news_item)

    news = {
        "Subject": subject,
        "News": main_news + also_featured_news
    }

    return news






